"""
Document Reader/Writer
By Ian Kruper

"""

import docx

def reader(doc): #Accepts document name, return string of text from document
    try:
        if doc[-3:] == "txt":
            with open(doc, 'r') as doc:
                text = doc.read()
                if len(text) < 20:
                    raise IOError
            return text
        elif doc[-4:] == "docx":
            word_doc = docx.Document(doc)
            text = ""
            for p in word_doc.paragraphs:
                text += p.text
                text += '\n'
            return text[0:-1]
        else:
            raise IOError
    except IOError:
        raise IOError
        

def writer(doc, text): #Accepts document name and text to write, returns nothing
    if doc[-3:] == "txt":
        with open(doc, 'w') as doc:
            doc.write(text)
    elif doc[-4:] == "docx":
        dest_doc = docx.Document()
        dest_doc.add_paragraph(text)
        dest_doc.save(doc)